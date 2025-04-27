```python
import os
import re
import logging
from docx import Document
from datetime import datetime
import pytz
import subprocess

logger = logging.getLogger(__name__)

def текущая_дата_лондон():
    return datetime.now(pytz.timezone("Europe/London")).strftime("%d.%m.%Y")

def очистить_имя_файла(text):
    return re.sub(r"[^\w\s-]", "", text, flags=re.UNICODE).strip()

def generate_pdf(путь_к_шаблону: str, текст: str) -> str:
    """
    Заполняет DOCX шаблон пользовательским текстом, конвертирует в PDF и возвращает путь к PDF.
    :param путь_к_шаблону: Путь к DOCX шаблону
    :param текст: Имя клиента
    :return: Путь к сгенерированному PDF
    """
    if not os.path.exists(путь_к_шаблону):
        logger.error(f"Шаблон не найден: {путь_к_шаблону}")
        raise FileNotFoundError(f"Шаблон не найден: {путь_к_шаблону}")

    logger.info(f"Генерация PDF с шаблоном: {путь_к_шаблону}, клиент: {текст}")
    дата = текущая_дата_лондон()
    имя_файла = очистить_имя_файла(текст) or "результат"
    путь_к_временному_файлу = f"{имя_файла}.docx"
    путь_к_выходному_файлу = f"{имя_файла}.pdf"

    # Заполнение DOCX шаблона
    doc = Document(путь_к_шаблону)

    for параграф in doc.paragraphs:
        if "{CLIENT}" in параграф.text:
            параграф.text = параграф.text.replace("{CLIENT}", текст)
        if "{DATE}" in параграф.text:
            параграф.text = параграф.text.replace("{DATE}", дата)

    # Замена в таблицах
    for таблица in doc.tables:
        for строка in таблица.rows:
            for ячейка in строка.cells:
                if "{CLIENT}" in ячейка.text:
                    ячейка.text = ячейка.text.replace("{CLIENT}", текст)
                if "{DATE}" in ячейка.text:
                    ячейка.text = ячейка.text.replace("{DATE}", дата)

    doc.save(путь_к_временному_файлу)
    logger.info(f"Сохранен временный DOCX: {путь_к_временному_файлу}")

    # Конвертация DOCX в PDF с использованием LibreOffice
    try:
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            ".",
            путь_к_временному_файлу
        ], check=True, capture_output=True, text=True)
        logger.info(f"Успешно создан PDF: {путь_к_выходному_файлу}")
    except subprocess.CalledProcessError as e:
        logger.error(f"Ошибка конвертации в PDF: {e.stderr}")
        raise Exception(f"Ошибка конвертации в PDF: {e.stderr}")

    # Удаление временного DOCX файла
    if os.path.exists(путь_к_временному_файлу):
        os.remove(путь_к_временному_файлу)
        logger.info(f"Удален временный DOCX: {путь_к_временному_файлу}")

    if not os.path.exists(путь_к_выходному_файлу):
        logger.error(f"PDF не создан: {путь_к_выходному_файлу}")
        raise FileNotFoundError(f"PDF не создан: {путь_к_выходному_файлу}")

    return путь_к_выходному_файлу
```
