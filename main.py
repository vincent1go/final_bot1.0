import os
import uuid
import subprocess
from datetime import datetime
from zoneinfo import ZoneInfo
import docx
import telegram
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application,
    CommandHandler,
    ContextTypes,
    ConversationHandler,
    MessageHandler,
    filters,
    CallbackQueryHandler,
)
import sqlite3
import logging
from dateutil.parser import parse

# Настройка логирования
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s", level=logging.INFO
)
logger = logging.getLogger(__name__)

# Состояния бота
SELECT_TEMPLATE, INPUT_NAME, CHANGE_DATE, INPUT_NEW_DATE, VIEW_BOOKMARKS, GENERATE_ANOTHER = range(6)

# Настройка базы данных
def init_db():
    conn = sqlite3.connect("bookmarks.db")
    c = conn.cursor()
    c.execute(
        """CREATE TABLE IF NOT EXISTS bookmarks
                 (user_id INTEGER, client_name TEXT, template_name TEXT, date TEXT)"""
    )
    conn.commit()
    conn.close()

# Инициализация базы данных
init_db()

# Сопоставление шаблонов
TEMPLATES = {
    "ur_recruitment": "template_ur.docx",
    "small_world": "template_small_world.docx",
    "imperative": "template_imperative.docx",
}

def replace_client_and_date(doc_path, client_name, date_str, template_key):
    try:
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"Шаблон {doc_path} не найден")
        
        doc = docx.Document(doc_path)
        
        # Замена Client
        client_replaced = False
        for para in doc.paragraphs:
            if "Client:" in para.text:
                if template_key == "small_world":
                    # Очистка строки перед "Client:" для Small World
                    para.text = f"Client: {client_name}"
                else:
                    para.text = para.text.replace("Client:", f"Client: {client_name}")
                client_replaced = True
                break
        if not client_replaced:
            logger.warning(f"Поле 'Client:' не найдено в {doc_path}")
        
        # Замена Date (дважды на последней странице)
        date_replaced_count = 0
        last_page_paragraphs = []
        current_page = []
        
        # Собираем абзацы, предполагая, что последняя страница — это последние абзацы
        for para in doc.paragraphs:
            current_page.append(para)
        last_page_paragraphs = current_page
        
        # Ищем "Date:" или "DATE:" в последних абзацах и заменяем дважды
        for para in last_page_paragraphs:
            if ("Date:" in para.text or "DATE:" in para.text) and date_replaced_count < 2:
                para.text = para.text.replace("Date:", f"Date: {date_str}")
                para.text = para.text.replace("DATE:", f"Date: {date_str}")
                date_replaced_count += 1
        if date_replaced_count != 2:
            logger.warning(f"Ожидалось 2 замены даты, выполнено {date_replaced_count} в {doc_path}")
        
        # Сохранение измененного документа
        temp_path = f"temp_{uuid.uuid4()}.docx"
        doc.save(temp_path)
        logger.info(f"Создан временный файл: {temp_path}")
        return temp_path
    except Exception as e:
        logger.error(f"Ошибка при обработке документа {doc_path}: {e}")
        raise

def convert_to_pdf(doc_path, client_name):
    pdf_path = f"{client_name}.pdf"
    try:
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"Временный файл {doc_path} не найден")
        
        # Вызов libreoffice для конвертации с ускорением
        logger.info(f"Запуск конвертации {doc_path} в PDF")
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--nofirststartwizard",  # Ускорение запуска LibreOffice
                "--convert-to",
                "pdf",
                "--outdir",
                os.path.dirname(doc_path) or ".",  # Убедимся, что outdir не пустой
                doc_path
            ],
            check=True,
            timeout=60  # Увеличенный таймаут
        )
        # Переименование файла
        temp_pdf = os.path.splitext(doc_path)[0] + ".pdf"
        if not os.path.exists(temp_pdf):
            raise FileNotFoundError(f"PDF-файл {temp_pdf} не создан")
        
        os.rename(temp_pdf, pdf_path)
        logger.info(f"PDF создан: {pdf_path}")
        return pdf_path
    except subprocess.CalledProcessError as e:
        logger.error(f"Ошибка конвертации в PDF: {e}")
        raise
    except FileNotFoundError as e:
        logger.error(f"Файл не найден: {e}")
        raise
    except subprocess.TimeoutExpired:
        logger.error("Превышено время ожидания для конвертации LibreOffice")
        raise
    except Exception as e:
        logger.error(f"Неизвестная ошибка при конвертации: {e}")
        raise

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("UR Recruitment", callback_data="ur_recruitment")],
        [InlineKeyboardButton("Small World", callback_data="small_world")],
        [InlineKeyboardButton("Imperative", callback_data="imperative")],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        "Добро пожаловать! Выберите шаблон:",
        reply_markup=reply_markup
    )
    return SELECT_TEMPLATE

async def select_template(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    
    template_key = query.data
    context.user_data["template_key"] = template_key
    
    await query.message.reply_text("Введите имя клиента:")
    return INPUT_NAME

async def receive_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    client_name = update.message.text.strip()
    context.user_data["client_name"] = client_name
    
    template_key = context.user_data["template_key"]
    template_path = os.path.join("templates", TEMPLATES[template_key])
    
    # Получение текущей даты в Киеве
    kyiv_tz = ZoneInfo("Europe/Kiev")
    current_date = datetime.now(kyiv_tz).strftime("%Y-%m-%d")
    context.user_data["date"] = current_date
    
    try:
        # Сообщение о начале генерации
        await update.message.reply_text("Ожидайте, ваш документ генерируется...")
        
        # Обработка документа
        temp_doc = replace_client_and_date(template_path, client_name, current_date, template_key)
        pdf_path = convert_to_pdf(temp_doc, client_name)
        
        # Отправка PDF
        with open(pdf_path, "rb") as f:
            await update.message.reply_document(document=f, filename=f"{client_name}.pdf")
        
        # Очистка временных файлов
        os.remove(temp_doc)
        os.remove(pdf_path)
        logger.info(f"Временные файлы удалены: {temp_doc}, {pdf_path}")
        
        # Предложение добавить в закладки, изменить дату или сгенерировать новый документ
        keyboard = [
            [InlineKeyboardButton("Добавить в закладки", callback_data="bookmark")],
            [InlineKeyboardButton("Изменить дату", callback_data="change_date")],
            [InlineKeyboardButton("Сгенерировать ещё один", callback_data="generate_another")],
            [InlineKeyboardButton("Начать заново", callback_data="start_over")],
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            "Документ сгенерирован! Что хотите сделать дальше?",
            reply_markup=reply_markup
        )
        return CHANGE_DATE
    except Exception as e:
        logger.error(f"Ошибка в receive_name: {e}")
        await update.message.reply_text(
            "Произошла ошибка при создании документа. Попробуйте снова или свяжитесь с поддержкой."
        )
        return ConversationHandler.END

async def bookmark(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    
    user_id = query.from_user.id
    client_name = context.user_data["client_name"]
    template_key = context.user_data["template_key"]
    date = context.user_data["date"]
    
    try:
        conn = sqlite3.connect("bookmarks.db")
        c = conn.cursor()
        c.execute(
            "INSERT INTO bookmarks (user_id, client_name, template_name, date) VALUES (?, ?, ?, ?)",
            (user_id, client_name, template_key, date)
        )
        conn.commit()
        conn.close()
        await query.message.reply_text("Документ успешно добавлен в закладки!")
    except Exception as e:
        logger.error(f"Ошибка при добавлении закладки: {e}")
        await query.message.reply_text("Ошибка при сохранении закладки. Попробуйте снова.")
    
    return CHANGE_DATE

async def change_date(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    
    await query.message.reply_text("Введите новую дату (например, 2025-04-28, 28.04.2025, 28/04/2025 и т.д.):")
    return INPUT_NEW_DATE

async def receive_new_date(update: Update, context: ContextTypes.DEFAULT_TYPE):
    new_date_input = update.message.text.strip()
    try:
        # Парсинг даты в любом формате
        parsed_date = parse(new_date_input)
        new_date = parsed_date.strftime("%Y-%m-%d")
    except ValueError:
        await update.message.reply_text("Неверный формат даты. Попробуйте снова (например, 2025-04-28, 28.04.2025):")
        return INPUT_NEW_DATE
    
    context.user_data["date"] = new_date
    client_name = context.user_data["client_name"]
    template_key = context.user_data["template_key"]
    template_path = os.path.join("templates", TEMPLATES[template_key])
    
    try:
        # Сообщение о начале генерации
        await update.message.reply_text("Ожидайте, ваш документ генерируется...")
        
        # Обработка документа с новой датой
        temp_doc = replace_client_and_date(template_path, client_name, new_date, template_key)
        pdf_path = convert_to_pdf(temp_doc, client_name)
        
        # Отправка PDF
        with open(pdf_path, "rb") as f:
            await update.message.reply_document(document=f, filename=f"{client_name}.pdf")
        
        # Очистка временных файлов
        os.remove(temp_doc)
        os.remove(pdf_path)
        logger.info(f"Временные файлы удалены: {temp_doc}, {pdf_path}")
        
        # Предложение вариантов
        keyboard = [
            [InlineKeyboardButton("Добавить в закладки", callback_data="bookmark")],
            [InlineKeyboardButton("Изменить дату снова", callback_data="change_date")],
            [InlineKeyboardButton("Сгенерировать ещё один", callback_data="generate_another")],
            [InlineKeyboardButton("Начать заново", callback_data="start_over")],
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            "Документ обновлен с новой датой! Что хотите сделать дальше?",
            reply_markup=reply_markup
        )
        return CHANGE_DATE
    except Exception as e:
        logger.error(f"Ошибка в receive_new_date: {e}")
        await update.message.reply_text(
            "Произошла ошибка при создании документа. Попробуйте снова или свяжитесь с поддержкой."
        )
        return ConversationHandler.END

async def generate_another(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    
    await query.message.reply_text("Введите имя нового клиента:")
    return GENERATE_ANOTHER

async def receive_another_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    client_name = update.message.text.strip()
    context.user_data["client_name"] = client_name
    
    template_key = context.user_data["template_key"]
    date = context.user_data["date"]
    template_path = os.path.join("templates", TEMPLATES[template_key])
    
    try:
        # Сообщение о начале генерации
        await update.message.reply_text("Ожидайте, ваш документ генерируется...")
        
        # Обработка документа
        temp_doc = replace_client_and_date(template_path, client_name, date, template_key)
        pdf_path = convert_to_pdf(temp_doc, client_name)
        
        # Отправка PDF
        with open(pdf_path, "rb") as f:
            await update.message.reply_document(document=f, filename=f"{client_name}.pdf")
        
        # Очистка временных файлов
        os.remove(temp_doc)
        os.remove(pdf_path)
        logger.info(f"Временные файлы удалены: {temp_doc}, {pdf_path}")
        
        # Предложение вариантов
        keyboard = [
            [InlineKeyboardButton("Добавить в закладки", callback_data="bookmark")],
            [InlineKeyboardButton("Изменить дату", callback_data="change_date")],
            [InlineKeyboardButton("Сгенерировать ещё один", callback_data="generate_another")],
            [InlineKeyboardButton("Начать заново", callback_data="start_over")],
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            "Документ сгенерирован! Что хотите сделать дальше?",
            reply_markup=reply_markup
        )
        return CHANGE_DATE
    except Exception as e:
        logger.error(f"Ошибка в receive_another_name: {e}")
        await update.message.reply_text(
            "Произошла ошибка при создании документа. Попробуйте снова или свяжитесь с поддержкой."
        )
        return ConversationHandler.END

async def start_over(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    
    context.user_data.clear()
    keyboard = [
        [InlineKeyboardButton("UR Recruitment", callback_data="ur_recruitment")],
        [InlineKeyboardButton("Small World", callback_data="small_world")],
        [InlineKeyboardButton("Imperative", callback_data="imperative")],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.message.reply_text(
        "Выберите шаблон:",
        reply_markup=reply_markup
    )
    return SELECT_TEMPLATE

async def view_bookmarks(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id
    try:
        conn = sqlite3.connect("bookmarks.db")
        c = conn.cursor()
        c.execute(
            "SELECT client_name, template_name, date FROM bookmarks WHERE user_id = ?",
            (user_id,)
        )
        bookmarks = c.fetchall()
        conn.close()
        
        if not bookmarks:
            await update.message.reply_text("У вас нет сохраненных закладок.")
            return ConversationHandler.END
        
        keyboard = [
            [
                InlineKeyboardButton(
                    f"{client_name} ({template_name}, {date})",
                    callback_data=f"bookmark_{client_name}_{template_name}_{date}"
                )
            ]
            for client_name, template_name, date in bookmarks
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(
            "Выберите сохраненный документ для повторной генерации:",
            reply_markup=reply_markup
        )
        return VIEW_BOOKMARKS
    except Exception as e:
        logger.error(f"Ошибка при просмотре закладок: {e}")
        await update.message.reply_text("Ошибка при загрузке закладок. Попробуйте снова.")
        return ConversationHandler.END

async def regenerate_bookmark(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    
    _, client_name, template_key, date = query.data.split("_", 3)
    context.user_data["client_name"] = client_name
    context.user_data["template_key"] = template_key
    context.user_data["date"] = date
    
    template_path = os.path.join("templates", TEMPLATES[template_key])
    
    try:
        # Сообщение о начале генерации
        await query.message.reply_text("Ожидайте, ваш документ генерируется...")
        
        # Обработка документа
        temp_doc = replace_client_and_date(template_path, client_name, date, template_key)
        pdf_path = convert_to_pdf(temp_doc, client_name)
        
        # Отправка PDF
        with open(pdf_path, "rb") as f:
            await query.message.reply_document(document=f, filename=f"{client_name}.pdf")
        
        # Очистка временных файлов
        os.remove(temp_doc)
        os.remove(pdf_path)
        logger.info(f"Временные файлы удалены: {temp_doc}, {pdf_path}")
        
        # Предложение вариантов
        keyboard = [
            [InlineKeyboardButton("Добавить в закладки", callback_data="bookmark")],
            [InlineKeyboardButton("Изменить дату", callback_data="change_date")],
            [InlineKeyboardButton("Сгенерировать ещё один", callback_data="generate_another")],
            [InlineKeyboardButton("Начать заново", callback_data="start_over")],
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await query.message.reply_text(
            "Документ повторно сгенерирован! Что хотите сделать дальше?",
            reply_markup=reply_markup
        )
        return CHANGE_DATE
    except Exception as e:
        logger.error(f"Ошибка в regenerate_bookmark: {e}")
        await query.message.reply_text(
            "Произошла ошибка при создании документа. Попробуйте снова или свяжитесь с поддержкой."
        )
        return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Операция отменена.")
    context.user_data.clear()
    return ConversationHandler.END

async def error_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logger.error(f"Update {update} caused error {context.error}")
    if update:
        if update.message:
            await update.message.reply_text(
                "Произошла ошибка. Попробуйте снова или свяжитесь с поддержкой."
            )
        elif update.callback_query:
            await update.callback_query.message.reply_text(
                "Произошла ошибка. Попробуйте снова или свяжитесь с поддержкой."
            )

def main():
    try:
        application = (
            Application.builder()
            .token("7677140739:AAGJcf8uhIKVdY44jqDKKlRM84_4_ndlrps")
            .build()
        )
        
        conv_handler = ConversationHandler(
            entry_points=[CommandHandler("start", start), CommandHandler("bookmarks", view_bookmarks)],
            states={
                SELECT_TEMPLATE: [CallbackQueryHandler(select_template)],
                INPUT_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, receive_name)],
                CHANGE_DATE: [
                    CallbackQueryHandler(bookmark, pattern="bookmark"),
                    CallbackQueryHandler(change_date, pattern="change_date"),
                    CallbackQueryHandler(generate_another, pattern="generate_another"),
                    CallbackQueryHandler(start_over, pattern="start_over"),
                ],
                INPUT_NEW_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, receive_new_date)],
                GENERATE_ANOTHER: [MessageHandler(filters.TEXT & ~filters.COMMAND, receive_another_name)],
                VIEW_BOOKMARKS: [CallbackQueryHandler(regenerate_bookmark, pattern="bookmark_.*")],
            },
            fallbacks=[CommandHandler("cancel", cancel)],
        )
        
        application.add_handler(conv_handler)
        application.add_error_handler(error_handler)
        
        # Проверка директории templates
        if not os.path.exists("templates"):
            logger.error("Директория templates не найдена")
            raise FileNotFoundError("Директория templates не найдена")
        
        # Запуск бота с вебхуком
        logger.info("Запуск приложения с вебхуком")
        application.run_webhook(
            listen="0.0.0.0",
            port=int(os.environ.get("PORT", 8443)),
            url_path="/webhook",
            webhook_url=f"https://{os.environ.get('RENDER_EXTERNAL_HOSTNAME')}/webhook"
        )
    except Exception as e:
        logger.error(f"Ошибка при запуске приложения: {e}")
        raise

if __name__ == "__main__":
    main()
